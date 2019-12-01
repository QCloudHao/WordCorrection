#!/usr/bin/env python3
# -*- coding: utf-8 -*-
__author__ = 'qyh'
__date__ = '2019/12/1 13:09'

from docx import Document
from nltk import  sent_tokenize, word_tokenize
from spelling_correcter import correct_text_generic
from docx.shared import RGBColor

# 文档中修改的单词个数
COUNT_CORRECT = 0

punkt_list = r",.?\"'!()/\\-<>:@#$%^&*~"

# word文档句柄
document = Document()


def write_correct_paragraph(file, paragraph_index):
    global COUNT_CORRECT

    # 每一段的内容
    paragraph = file.paragraphs[paragraph_index].text.strip()
    # 进行句子的划分
    sentences = sent_tokenize(text=paragraph)
    # 词语划分
    words_list = [word_tokenize(sentence) for sentence in sentences]

    p = document.add_paragraph(' ' * 7)  # 段落句柄

    for word_list in words_list:
        for word in word_list:
            # 每一句话第一个单词的第一个字母大写，并空两格
            if word_list.index(word) == 0 and words_list.index(word_list) == 0:
                if word not in punkt_list:
                    p.add_run(' ')
                    # 修改单词，如果单词正确，则返回原单词
                    correct_word = correct_text_generic(word)
                    # 如果该单词有修改，则颜色为红色
                    if correct_word != word:
                        colored_word = p.add_run(correct_word[0].upper() + correct_word[1:])
                        font = colored_word.font
                        font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word[0].upper() + correct_word[1:])
                else:
                    p.add_run(word)
            else:
                p.add_run(' ')
                # 修改单词，如果单词正确，则返回原单词
                correct_word = correct_text_generic(word)
                if word not in punkt_list:
                    # 如果该单词有修改，则颜色为红色
                    if correct_word != word:
                        colored_word = p.add_run(correct_word)
                        font = colored_word.font
                        font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word)
                else:
                    p.add_run(word)


if __name__ == '__main__':
    file = Document("./Spelling Error.docx")
    for i in range(len(file.paragraphs)):
        write_correct_paragraph(file, i)
    document.save('./correct_document.docx')

    print('修改并保存文件完毕！')
    print('一共修改了%d处。' % COUNT_CORRECT)
